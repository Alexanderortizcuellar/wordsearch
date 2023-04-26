import random
import string
import csv
import time
import datetime
import typing
from openpyxl import Workbook
from openpyxl.worksheet.worksheet import Worksheet
from openpyxl.cell.cell import Cell
from openpyxl.styles import Font, Border, Alignment, Side, PatternFill
from openpyxl.utils import get_column_letter
import json
from rich.console import Console
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter, legal
from reportlab.platypus import Table, TableStyle, Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.units import mm
from reportlab.lib.styles import ParagraphStyle, TA_CENTER
from more_itertools import grouper
from jinja2.environment import Environment
from jinja2.loaders import FileSystemLoader
from docx import Document
from unidecode import unidecode


def clean_words(words: list[str]):
    words = [unidecode(word) for word in words]
    words = sorted(words, key=len, reverse=True)
    words = list(map(lambda x: x.upper().strip(), words))
    words = [word.replace('"', "") for word in words]
    return words


color_names = [
    "orange4", "bright_blue", "blue", "yellow",
    "green", "magenta", "red", "cyan", "dark_orange",
    "steel_blue3", "bright_red", "dark_red"
]


list_colors = [
    colors.darkgreen, colors.blue, colors.red,
    colors.darkorange, colors.brown, colors.darkblue,
    colors.darkgoldenrod]


class WordSearch():
    def __init__(self, words: list, width: int, height, show_answers=False) -> None:
        self.words = words
        self.width = width
        self.height = height
        self.directions = [(1, 0), (0, 1), (1, 1)]
        self.show_answers = show_answers
        self.generate()

    def process_words(self):
        """ processing the words to optimize
                the code
         """
        # sorting the words by length
        words = sorted(self.words, key=len, reverse=True)
        # making all the words upper case
        words = list(map(lambda x: x.upper().strip(), words))

        return words

    def make_grid(self):
        grid = [["*" for col in range(self.width)]
                for row in range(self.height)]
        return grid

    def place_word(self, positions, grid):
        for i in range(len(positions)):
            grid[positions[i][0]][positions[i][1]] = positions[i][2]

    def fill(self, grid):
        for row in range(self.height):
            for item in range(self.width):
                if grid[row][item] == "*":
                    grid[row][item] = random.choice(string.ascii_uppercase)

    def retry(self, word, grid: list):
        candidates = []
        for row in range(self.height):
            for item in range(self.width):
                if grid[row][item] == "*" or grid[row][item] == word[0]:
                    candidates.append((row, item))
        for p in candidates:
            done = False
            for d in self.directions:
                positions = []
                found = False
                for i, c in enumerate(word):
                    if p[0]+i*d[0] >= self.width or p[1]+i*d[1] >= self.height:
                        positions.clear()
                        break
                    pos = grid[p[0]+i*d[0]][p[1]+i*d[1]]
                    if pos == "*" or pos == c:
                        positions.append((p[0]+i*d[0], p[1]+i*d[1], c))
                    else:
                        positions.clear()
                        break
                else:
                    self.place_word(positions, grid)
                    found = True
                    return (True, positions)
                    break
                if found == True:
                    break
                continue
            if found == True:
                break
        return False, positions

    def locate(self, word, grid):
        tries = 0
        placed = False
        while placed == False and tries <= 2000:
            random.shuffle(self.directions)
            d = random.choice(self.directions)
            word = random.choice([word, word[::-1]])
            xstart = random.randint(0, self.width)
            ystart = random.randint(0, self.height)
            xstart = xstart if xstart + \
                len(word) < self.width else xstart-len(word)
            ystart = ystart if ystart + \
                len(word) < self.height else ystart-len(word)
            positions = []
            for i, c in enumerate(word):
                if xstart+i*d[0] < 0 or ystart+i*d[1] < 0:
                    break
                if xstart+i*d[0] >= self.width or ystart+i*d[1] >= self.height:
                    break
                pos = grid[xstart+i*d[0]][ystart+i*d[1]]
                if pos == "*" or pos == word[i]:
                    #grid[xstart+i*d[0]][ystart+i*d[1]] = word[i]
                    position = xstart+i*d[0], ystart+i*d[1], c
                    positions.append(position)
                else:
                    positions.clear()
                    break
            else:
                self.place_word(positions, grid)
                return positions
                placed = True
            tries += 1
            if tries == 2000:
                added, positions = self.retry(word, grid)
                if not added:
                    print(
                        f"tried {tries} times and could not add {word}, skipping")
                return positions
            continue

    def predict_width_height(self, words, ratio=1.3):
        words = clean_words(words)
        if len(words) > 0:
            len_longest = len(words[0])
            letters = sum(list(map(len, words)))
            total = letters*2.3
            dimension = round(total**0.5)
            if dimension < len_longest:
                dimension = len_longest
            return dimension

    def color_word(self, word):
        color = random.choice(color_names)
        word_position = self.all_positions[word]
        for pos in word_position:
            self.grid[pos[0]][pos[1]
                              ] = f"[bold {color}]{pos[2]}[/bold {color}]"

    def pos_to_style(self, pos: list):
        color = random.choice(list_colors)
        pos = tuple(tuple(reversed(p[0:2])) for p in pos)
        if len(pos) > 0:
            start = pos[0]
            end = pos[-1]
            if start[0] == end[0] or start[1] == end[1]:
                style = ("BACKGROUND", start, end, color)
            else:
                style = [("BACKGROUND", p, p, color) for p in pos]
            return style

    def export_pdf(self, filename: str, title: str, answers=False, fontsize=12):
        pdf = SimpleDocTemplate(filename)
        table = Table(self.grid, colWidths=6*mm, rowHeights=6*mm)
        STYLE = []
        for p in self.all_positions:
            styles = self.pos_to_style(self.all_positions[p])
            if isinstance(styles, list):
                for s in styles:
                    STYLE.append(s)
            else:
                if styles is not None:
                    STYLE.append(styles)
        border = ("BOX", (0, 0), (-1, -1), 1, colors.black)
        STYLE.append(border)
        STYLE.append(('ALIGN', (0, 0), (-1, -1), 'CENTRE'))
        STYLE.append(('FONTSIZE', (0, 0), (-1, -1), fontsize))
        table_style = [border] if answers == False else STYLE
        table.setStyle(table_style)
        paragraph_style = ParagraphStyle(
            name="p", alignment=TA_CENTER, fontSize=18, spaceAfter=30)
        paragraph = Paragraph(title, style=paragraph_style)

        d = list(grouper(self.words, 4))
        words_table = Table(d)
        pdf.build([paragraph, table, Spacer(pdf.width, 10*mm), words_table])

    def show(self, answer: bool = False):
        console = Console()
        if answer == True:
            for word in self.process_words():
                self.color_word(word)
        for row in self.grid:
            console.print(" ".join(row))

    def generate(self):
        self.all_positions = {}
        words = self.process_words()
        grid = self.make_grid()
        for word in words:
            positions = self.locate(word, grid)
            self.all_positions[word] = positions
        if not self.show_answers:
            self.fill(grid)
        self.grid = grid
        placed = [word for word in self.all_positions if len(
            self.all_positions[word]) >= 1]
        not_placed = [word for word in self.all_positions if len(
            self.all_positions[word]) == 0]
        number_letters_words = sum([len(key) for key in self.all_positions])
        self.output = {
            "wordsearch": grid, "positions": self.all_positions,
            "words": list(self.all_positions.keys()),
            "number-words": len(self.all_positions),
            "number-words-placed": len(placed),
            "number-words-no-placed": len(not_placed),
            "words-placed": placed, "words-not-placed": not_placed,
            "number-letters": self.width*self.height,
            "number-letter-words": number_letters_words,
            "number-random-letters": (self.width*self.height) - number_letters_words,
            "time-created": datetime.datetime.now().strftime("%Y-%m-%d-%I:%M:%S %p")
        }

    def export_text(self, filename: str):
        with open(filename, "w") as file:
            for row in self.grid:
                file.write(f'{" ".join(row)}\n')
            file.write("________________________________\n")
            for word in self.all_positions.keys():
                file.write(word+"\n")

    def export_csv(self, filename: str, delimiter=",", encoding="UTF-8"):
        with open(filename, "w", newline="", encoding=encoding) as file:
            writer = csv.writer(file, delimiter=delimiter)
            writer.writerows(self.grid)
            writer.writerow(["" for x in range(self.width)])
            for w in self.all_positions.keys():
                writer.writerow([w])

    def pos_to_excel(self, pos: list):
        "due to the grid does not start from the first row and column"
        new_pos = [(f"{get_column_letter(p[1]+2)}{p[0]+3}", p[2]) for p in pos]
        return new_pos

    def style_ws(self, ws: Worksheet, coords: dict):
        colors = ["AA00FF00", "AAFF0000", "AAFFFF00", "AA0000FF"]
        for key, val in list(coords.items()):
            color = random.choice(colors)
            for pos in coords[key]:
                ws[pos[0]].fill = PatternFill(
                    start_color=color, end_color=color, fill_type="solid")

    def all_pos_to_excel_coords(self):
        self.excel_cords = {}
        for key, val in list(self.all_positions.items()):
            self.excel_cords[key] = self.pos_to_excel(val)

    def export_excel(self, filename: str, sheetname: str = "wordsearch", answers=False, fontname="Arial", fontsize=16):
        self.all_pos_to_excel_coords()
        wb = Workbook()
        ws: Worksheet = wb.active
        ws["A2"] = ""
        end_letter = get_column_letter(self.width+1)
        ws.merge_cells(f"B1:{end_letter}1")

        values = self.grid
        for row in values:
            ws.append(row)

        for col in range(2, self.width+2):
            letter = get_column_letter(col)
            ws.column_dimensions[letter].width = 5
        ws.insert_cols(1, 1)
        ws.column_dimensions["A"].width = 1
        ws["B1"] = "wordsearch"
        ws["B1"].font = Font(name="Arial", size=21)
        ws["B1"].alignment = Alignment(horizontal="center", vertical="center")
        cell: Cell
        for row in ws.iter_rows(min_row=3, max_row=self.width+2, min_col=2, max_col=self.width+1):
            for cell in row:
                cell.font = Font(size=fontsize, name=fontname, bold=True)
                cell.alignment = Alignment(
                    horizontal="center", vertical="center")
                side = Side(border_style="thin", color='FF000000')
                cell.border = Border(left=side, right=side,
                                     top=side, bottom=side)
                cell.fill = PatternFill(
                    start_color="e1e1e1", end_color="a6a6a6", fill_type="solid")
        for row in ws.iter_rows(min_row=3, min_col=self.width+3, max_row=2+len(self.all_positions.keys()), max_col=self.width+3):
            for cell in row:
                cell.value = list(self.all_positions.keys())[cell.row-3]
        ws.column_dimensions[get_column_letter(self.width+3)].width = 18
        ws.title = sheetname if sheetname != "" else "wordsearch"
        if answers == True:
            self.style_ws(ws, self.excel_cords)
        wb.save(filename)

    def export_json(self, filename=str):
        wordsearch = {}
        wordsearch["Wordsearch"] = self.grid
        wordsearch["Positions"] = self.all_positions
        wordsearch["Words"] = list(self.all_positions.keys())
        with open(filename, "w", encoding="utf-8") as file:
            json.dump(wordsearch, file)

    def export_docx(self, filename, title="Wordsearch", answers=False):
        doc = Document()
        doc.add_heading(title, 0)
        table = doc.add_table(rows=self.height, columns=self.height)
        for row in self.grid:
            for col in self.grid[0]:
                table.rows[row].cells[col].text = col
        doc.save(filename)

    def export_html(self, filename):
        pass
